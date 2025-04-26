# document_writer.py
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def apply_base_formatting(run, format_info):
    """Apply base formatting (font name, size, color) to a run"""
    if format_info.get('font_name'):
        run.font.name = format_info['font_name']
    
    if format_info.get('font_size'):
        # Convert to points if it's not already
        font_size = format_info['font_size']
        if isinstance(font_size, (int, float)):
            run.font.size = Pt(font_size)
    
    if format_info.get('font_color'):
        # Convert from serialized format
        color_str = format_info['font_color']
        if color_str and ':' in color_str:
            try:
                r, g, b = map(int, color_str.split(':'))
                run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, TypeError):
                # If there's an error parsing the color, skip it
                pass

def apply_formatting_to_paragraph(paragraph, text, format_info):
    """Function to parse the formatted answer and apply proper Word formatting"""
    # Process for bullet points - manually create bullet effect without using styles
    if text.startswith('•'):
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        text = text[1:].strip()
        # Add a bullet character at the beginning
        run = paragraph.add_run('• ')
        apply_base_formatting(run, format_info)
    
    # Look for formatting tags
    remaining_text = text
    position = 0
    
    # If the text starts with a bullet, we've already processed it
    if text.startswith('•'):
        remaining_text = text[1:].strip()
    
    while position < len(remaining_text):
        # Look for the next formatting tag
        bold_start = remaining_text.find('<b>', position)
        italic_start = remaining_text.find('<i>', position)
        bold_italic_start = remaining_text.find('<bi>', position)
        
        # Find the earliest tag
        next_tag_position = min(
            bold_start if bold_start != -1 else float('inf'),
            italic_start if italic_start != -1 else float('inf'),
            bold_italic_start if bold_italic_start != -1 else float('inf')
        )
        
        if next_tag_position == float('inf'):
            # No more tags, add the remaining text as regular
            if position < len(remaining_text):
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
            break
        
        # Add text before the tag
        if next_tag_position > position:
            run = paragraph.add_run(remaining_text[position:next_tag_position])
            apply_base_formatting(run, format_info)
        
        # Process the tag
        if next_tag_position == bold_start:
            end_tag = remaining_text.find('</b>', bold_start)
            if end_tag != -1:
                content = remaining_text[bold_start+3:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.bold = True
                position = end_tag + 4  # Move past </b>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
                
        elif next_tag_position == italic_start:
            end_tag = remaining_text.find('</i>', italic_start)
            if end_tag != -1:
                content = remaining_text[italic_start+3:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.italic = True
                position = end_tag + 4  # Move past </i>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
                
        elif next_tag_position == bold_italic_start:
            end_tag = remaining_text.find('</bi>', bold_italic_start)
            if end_tag != -1:
                content = remaining_text[bold_italic_start+4:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.bold = True
                run.italic = True
                position = end_tag + 5  # Move past </bi>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
    
    # Set paragraph alignment if specified
    if format_info.get('alignment'):
        try:
            # Convert from string back to enum if needed
            alignment = format_info['alignment']
            if isinstance(alignment, str) and alignment.isdigit():
                paragraph.alignment = int(alignment)
        except (ValueError, TypeError):
            # If there's an error parsing the alignment, skip it
            pass

def create_table_from_description(doc, table_desc, format_info):
    """Process table formatting"""
    # Parse table description: row1col1|row1col2;row2col1|row2col2
    rows = table_desc.split(';')
    num_rows = len(rows)
    num_cols = len(rows[0].split('|'))
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_text in enumerate(rows):
        cells = row_text.split('|')
        for j, cell_text in enumerate(cells):
            if j < num_cols:  # Ensure we don't exceed column count
                cell = table.cell(i, j)
                paragraph = cell.paragraphs[0]
                apply_formatting_to_paragraph(paragraph, cell_text.strip(), format_info)
                
    return table

def insert_answers_in_document(doc, qa_data):
    """Insert answers into the document using the saved question-answer data"""
    # Extract data from the qa_data
    qa_pairs = qa_data.get("qa_pairs", [])
    question_indices = qa_data.get("question_indices", [])
    question_texts = qa_data.get("question_texts", [])
    question_formats = qa_data.get("question_formats", [])
    
    # Create a mapping of questions to answers
    qa_map = {qa["question"].strip(): qa["answer"] for qa in qa_pairs}
    
    modified = False
    insertions = 0
    
    # Process each question paragraph
    for idx, (question_text, format_info) in enumerate(zip(question_texts, question_formats)):
        question_index = question_indices[idx]
        
        # Find the matching answer
        matching_answer = None
        matching_question = None
        
        # Try to find an exact match first
        if question_text in qa_map:
            matching_answer = qa_map[question_text]
            matching_question = question_text
        else:
            # Try to find a fuzzy match
            for q in qa_map:
                # Check if the current question is similar to any in our map
                if question_text in q or q in question_text:
                    matching_answer = qa_map[q]
                    matching_question = q
                    break
        
        if matching_answer:
            # Force insert answer after the question
            print(f"Inserting answer after: {question_text[:50]}...")
            
            # Create a new paragraph for "Answer:" heading
            p_answer_heading = doc.add_paragraph()
            doc._body._body.insert(question_index + 1, p_answer_heading._p)
            p_answer_heading.style = 'Normal'  # Use Normal style which should exist in every document
            p_answer_heading.paragraph_format.space_before = Pt(12)
            
            # Apply the same formatting as the question to the "Answer:" label
            r_answer_heading = p_answer_heading.add_run("Answer: ")
            apply_base_formatting(r_answer_heading, format_info)
            r_answer_heading.bold = True
            r_answer_heading.font.color.rgb = RGBColor(0, 102, 0)  # Green color for "Answer:" text
            
            # Current insertion position
            current_pos = question_index + 2
            
            # Process answer text with appropriate formatting
            answer_lines = matching_answer.split('\n')
            for line in answer_lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check for table
                table_match = re.search(r'<table>(.*?)</table>', line)
                if table_match:
                    table_desc = table_match.group(1)
                    table = create_table_from_description(doc, table_desc, format_info)
                    doc._body._body.insert(current_pos, table._tbl)
                    current_pos += 1
                    continue
                    
                # Check for diagram placeholder
                diagram_match = re.search(r'<diagram>(.*?)</diagram>', line)
                if diagram_match:
                    p_diagram = doc.add_paragraph()
                    doc._body._body.insert(current_pos, p_diagram._p)
                    r_diagram = p_diagram.add_run("[Diagram: " + diagram_match.group(1) + "]")
                    apply_base_formatting(r_diagram, format_info)
                    r_diagram.italic = True
                    current_pos += 1
                    continue
                
                # Regular paragraph with potential formatting
                p_answer = doc.add_paragraph()
                doc._body._body.insert(current_pos, p_answer._p)
                p_answer.style = 'Normal'  # Use Normal style which should exist in every document
                
                # If not a bullet point, add indentation
                if not line.startswith('•'):
                    p_answer.paragraph_format.left_indent = Pt(20)
                
                # Apply formatting to the text
                apply_formatting_to_paragraph(p_answer, line, format_info)
                current_pos += 1
            
            # Add a blank line after the answer
            p_blank = doc.add_paragraph()
            doc._body._body.insert(current_pos, p_blank._p)
            current_pos += 1
            
            modified = True
            insertions += 1
            
            # Adjust subsequent indices to account for insertions
            offset = current_pos - (question_index + 1)
            for j in range(len(question_indices)):
                if question_indices[j] > question_index:
                    question_indices[j] += offset
    
    print(f"Inserted {insertions} answers into the document.")
    return modified

def write_answers_to_document(input_file, answers_json, output_file):
    """Insert answers from JSON into the document and save as a new file"""
    print(f"Loading document: {input_file}")
    print(f"Loading answers from: {answers_json}")
    
    # Load the document
    doc = Document(input_file)
    
    # Load the answers data
    with open(answers_json, 'r', encoding='utf-8') as f:
        qa_data = json.load(f)
    
    # Insert answers into the document
    modified = insert_answers_in_document(doc, qa_data)
    
    if modified:
        # Save the modified document
        doc.save(output_file)
        print(f"Successfully created document with answers: {output_file}")
        return True
    else:
        print("No modifications were made to the document.")
        return False

if __name__ == "__main__":
    # This allows the module to be run directly for testing
    input_file = "E:\\ClgStu\\Sem 5\\Computer Networks(3150710)\\3150710-Computer-Network Lab manual.docx"
    answers_json = "answers_data.json"
    output_file = "lab_manual_with_answers.docx"
    write_answers_to_document(input_file, answers_json, output_file)