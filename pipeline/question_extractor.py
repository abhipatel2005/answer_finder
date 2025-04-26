import re
import json
import os
from docx import Document
import fitz  # PyMuPDF for PDF extraction

def extract_lab_manual_text_docx(doc_path):
    """Extract full text from DOCX for context"""
    doc = Document(doc_path)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    
    return "\n".join(full_text)

def extract_lab_manual_text_pdf(pdf_path):
    """Extract text from PDF for context"""
    full_text = []
    
    try:
        # Open the PDF file
        doc = fitz.open(pdf_path)
        
        # Extract text from each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            full_text.append(page.get_text())
        
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def clean_text(text):
    """Clean up text formatting"""
    # Replace multiple whitespace with single space
    text = re.sub(r'\s+', ' ', text)
    # Replace bullet points with dashes
    text = re.sub(r'•', '- ', text)
    # Clean up any remaining special characters
    text = re.sub(r'[^\w\s\.\,\:\;\-\(\)\[\]\{\}\?\!]', '', text)
    return text.strip()

def extract_questions_from_quiz_section(quiz_section):
    """Extract numbered questions from a quiz section"""
    # Clean up the text first
    text = clean_text(quiz_section)
    
    # Regex to capture numbered questions followed by any text
    question_regex = r'\d+\.\s*(.*?)(?=\d+\.\s|$)'
    questions = re.findall(question_regex, text, re.DOTALL)
    
    # Clean up each question
    return [clean_text(q) for q in questions if len(clean_text(q)) > 10]

def extract_quiz_sections(full_text):
    """Extract quiz sections from the document text"""
    # Pattern to identify quiz sections
    quiz_sections_pattern = r'(Quiz\(.*?\)|Quiz: \(Sufficient space to be provided for the answers\))(.*?)(References used by the students:?|Suggested Reference:?)'
    quiz_sections = re.findall(quiz_sections_pattern, full_text, re.DOTALL)
    
    extracted_sections = []
    if quiz_sections:
        print(f"Found {len(quiz_sections)} quiz sections")
        for section in quiz_sections:
            quiz_content = section[1].strip()  # Content between 'Quiz' and the reference section
            extracted_sections.append(quiz_content)
    else:
        print("No quiz sections found using regular pattern")
    
    return extracted_sections

def find_question_paragraphs_docx(doc_path):
    """Find all paragraphs containing numbered questions and extract formatting info (DOCX only)"""
    doc = Document(doc_path)
    question_paragraphs = []
    question_texts = []
    question_formats = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Match questions that start with a number followed by period, then text
        match = re.match(r'^\s*(\d+)\.\s*(.*?)$', text)
        if match:
            number = match.group(1)
            question = match.group(2).strip()
            if question:  # Only add if there's actual question text
                question_paragraphs.append(i)
                question_texts.append(question)
                
                # Store font information
                format_info = {
                    'font_name': None,
                    'font_size': None,
                    'font_color': None,
                    'alignment': None
                }
                
                # Extract format from first run that contains text
                for run in para.runs:
                    if run.text.strip():
                        if run.font.name:
                            format_info['font_name'] = run.font.name
                        if run.font.size:
                            # Convert to points for serialization
                            if hasattr(run.font.size, 'pt'):
                                format_info['font_size'] = run.font.size.pt
                        if run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                            # Convert RGB color to hex string for serialization
                            rgb = run.font.color.rgb
                            if rgb:
                                format_info['font_color'] = f"{rgb[0]}:{rgb[1]}:{rgb[2]}"
                        break
                
                # Store paragraph alignment as string for serialization
                if para.alignment:
                    format_info['alignment'] = str(para.alignment)
                
                question_formats.append(format_info)
                print(f"Found question {number}: {question}")
    
    return question_paragraphs, question_texts, question_formats

def extract_all_questions(file_path):
    """Extract all questions from either PDF or DOCX document"""
    file_ext = os.path.splitext(file_path)[1].lower()
    all_questions = []
    full_text = ""
    question_indices = []
    question_texts = []
    question_formats = []
    
    # Extract text based on file type
    if file_ext == '.pdf':
        print(f"Processing PDF document: {file_path}")
        full_text = extract_lab_manual_text_pdf(file_path)
        quiz_sections = extract_quiz_sections(full_text)
        
        # Extract questions from quiz sections
        if quiz_sections:
            for i, section in enumerate(quiz_sections):
                questions = extract_questions_from_quiz_section(section)
                print(f"Quiz section {i+1}: Found {len(questions)} questions")
                all_questions.extend(questions)
        else:
            # If no quiz sections found, try to extract all numbered questions
            all_questions = extract_questions_from_quiz_section(full_text)
            print(f"Found {len(all_questions)} numbered questions in the document")
            
    elif file_ext == '.docx':
        print(f"Processing DOCX document: {file_path}")
        full_text = extract_lab_manual_text_docx(file_path)
        question_indices, question_texts, question_formats = find_question_paragraphs_docx(file_path)
        
        quiz_sections = extract_quiz_sections(full_text)
        
        # Extract questions from quiz sections
        if quiz_sections:
            for i, section in enumerate(quiz_sections):
                questions = extract_questions_from_quiz_section(section)
                print(f"Quiz section {i+1}: Found {len(questions)} questions")
                all_questions.extend(questions)
        else:
            # If no quiz sections, use the questions found in paragraphs
            all_questions = question_texts if question_texts else extract_questions_from_quiz_section(full_text)
            print(f"Found {len(all_questions)} numbered questions in the document")
    else:
        print(f"Unsupported file format: {file_ext}")
    
    return all_questions, full_text, question_indices, question_texts, question_formats

def create_student_answer(question, context):
    """Create a simple student-like answer based on the question type"""
    question = question.lower()
    
    # Extract potential topic from the question (for contextual answers)
    topics = ["hub", "router", "switch", "bridge", "gateway", "network", "protocol", "ethernet", "packet", "transmission", "connectivity"]
    question_topic = None
    for topic in topics:
        if topic in question:
            question_topic = topic
            break
    
    # Create answer based on question type
    if "what is" in question:
        if question_topic:
            return f"A {question_topic} is a device used in computer networks for managing data communication."
        return "This is a networking component that helps transmit data between devices."
        
    elif "difference between" in question or "differentiate" in question:
        match = re.search(r"(difference between|differentiate).*?(\w+).*?(\w+)", question)
        if match:
            term1 = match.group(2)
            term2 = match.group(3)
            return f"The main difference between {term1} and {term2} is in their functionality and which OSI layer they operate at."
        return "The key difference is in their functionality and network layers they operate on."
        
    elif "explain" in question or "describe" in question:
        if question_topic:
            return f"The {question_topic} works by routing data packets through the network according to specific rules."
        return "This concept functions by processing data according to network protocols and rules."
        
    elif "advantage" in question or "benefit" in question:
        if question_topic:
            return f"The main advantages of {question_topic} include faster data transmission and improved network efficiency."
        return "The advantages include better performance, improved security, and easier network management."
        
    elif "disadvantage" in question or "limitation" in question:
        if question_topic:
            return f"Some limitations of {question_topic} include higher cost and more complex configuration."
        return "Some disadvantages include increased complexity and potential compatibility issues."
        
    else:
        return "This relates to computer networking concepts covered in the lab manual."

# Function that matches the import in main.py
def extract_and_save_questions(input_file, output_json):
    """Extract questions from document and save to JSON file"""
    print(f"Processing lab manual: {input_file}")
    
    # Extract questions and text from the document
    all_questions, full_text, question_indices, question_texts, question_formats = extract_all_questions(input_file)
    
    if all_questions:
        print(f"Found a total of {len(all_questions)} questions")
        
        # Create the data structure to save
        question_data = {
            "questions": all_questions,
            "context": full_text,
            "question_indices": question_indices,
            "question_texts": question_texts,
            "question_formats": question_formats
        }
        
        # Save to JSON file
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump(question_data, f, ensure_ascii=False, indent=2)
        
        print(f"Successfully saved {len(all_questions)} questions to {output_json}")
        return True
    else:
        print("No questions found in the document.")
        return False

def extract_and_save_questions_and_answers(input_file, output_json):
    """Extract questions and answers from document and save to JSON file with improved formatting"""
    print(f"Processing lab manual: {input_file}")
    
    # Extract questions from the document
    all_questions, full_text, question_indices, question_texts, question_formats = extract_all_questions(input_file)
    
    # Generate simple, well-formatted student-style answers
    all_answers = []
    for question in all_questions:
        answer = create_student_answer(question, full_text)
        all_answers.append(answer)
    
    if all_questions:
        print(f"Found a total of {len(all_questions)} questions")
        
        # Create better formatted QnA pairs for output
        formatted_qna = []
        for i, (q, a) in enumerate(zip(all_questions, all_answers)):
            formatted_qna.append({
                "question_number": i + 1,
                "question_text": q.strip(),
                "answer_text": a.strip()
            })
        
        # Create the data structure to save
        data = {
            "total_questions": len(all_questions),
            "qna_pairs": formatted_qna,
            "context": full_text[:1000] + "..." if len(full_text) > 1000 else full_text  # Truncate for readability
        }
        
        # Save to JSON file
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"Successfully saved {len(all_questions)} questions and answers to {output_json}")
        return True
    else:
        print("No questions found in the document.")
        return False

def generate_pdf_ready_content(questions, answers):
    """Generate properly formatted content for PDF generation"""
    content = []
    
    for i, (question, answer) in enumerate(zip(questions, answers)):
        q_num = i + 1
        
        # Format question with proper spacing
        q_formatted = f"Question {q_num}: {question.strip()}"
        
        # Format answer with proper spacing and indentation
        a_formatted = f"Answer {q_num}: {answer.strip()}"
        
        # Add to content with proper spacing between QnA pairs
        content.append({
            "question": q_formatted,
            "answer": a_formatted
        })
    
    return content

if __name__ == "__main__":
    # Example usage
    input_file = "lab_manual.pdf"  # Can be either PDF or DOCX
    output_json = "questions_data.json"
    
    # For just questions
    extract_and_save_questions(input_file, output_json)
    
    # For questions and answers with improved formatting
    extract_and_save_questions_and_answers(input_file, "questions_with_answers.json")
    
    # Generate PDF-ready content
    all_questions, _, _, _, _ = extract_all_questions(input_file)
    if all_questions:
        all_answers = [create_student_answer(q, "") for q in all_questions]
        pdf_content = generate_pdf_ready_content(all_questions, all_answers)
        
        with open("pdf_ready_content.json", 'w', encoding='utf-8') as f:
            json.dump(pdf_content, f, ensure_ascii=False, indent=2)