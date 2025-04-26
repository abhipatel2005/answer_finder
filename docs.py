# import os
# import re
# import asyncio
# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import google.generativeai as genai
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()

# # Initialize Gemini API
# GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# genai.configure(api_key=GEMINI_API_KEY)

# # Extract full text from DOCX for context
# def extract_lab_manual_text(doc):
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     full_text.append(para.text)
#     return "\n".join(full_text)

# # Process quiz questions using Gemini API with detailed answers
# async def generate_detailed_answer(question, context):
#     model = genai.GenerativeModel('gemini-2.0-flash')
    
#     prompt = f"""You are an expert tutor writing content for a lab manual. 
#     Please generate a detailed, educational answer to the following question based on the lab manual content.
    
#     Format your answer with appropriate structure:
#     - Use bullet points for listing items or steps
#     - Use bold for important terms or concepts
#     - Use italics for emphasis or definitions
#     - Create tables if data needs to be compared
#     - Add diagrams or explanations where needed
    
#     Make it look like a professional textbook answer that's educational and comprehensive.
    
#     Question: {question}
    
#     Context: {context[:10000]}"""
    
#     try:
#         response = model.generate_content(prompt)
#         answer = response.text.strip() if response.text else 'No answer generated.'
#         print(f"\nQ: {question}\nA: {answer}\n")
#         return answer
#     except Exception as err:
#         print(f"Error generating answer for '{question}': {str(err)}")
#         return 'Error generating answer.'

# # Improved Question Extraction Regex
# def extract_questions_from_quiz_section(quiz_section):
#     # Regex to capture numbered questions followed by any text, even if it doesn't end in a '?' or '.'.
#     question_regex = r'\d+\.\s*(.*?)(?=\d+\.\s|$)'
#     questions = re.findall(question_regex, quiz_section, re.DOTALL)
#     return [q.strip() for q in questions]

# # Function to find all paragraphs containing numbered questions
# def find_question_paragraphs(doc):
#     question_paragraphs = []
#     question_texts = []
    
#     for i, para in enumerate(doc.paragraphs):
#         text = para.text.strip()
#         # Match questions that start with a number followed by period, then text
#         match = re.match(r'^\s*(\d+)\.\s*(.*?)$', text)
#         if match:
#             number = match.group(1)
#             question = match.group(2).strip()
#             if question:  # Only add if there's actual question text
#                 question_paragraphs.append(i)
#                 question_texts.append(question)
#                 print(f"Found question {number}: {question}")
    
#     return question_paragraphs, question_texts

# # Insert answers after questions in the original document - FORCED VERSION
# def insert_answers_in_document_forced(doc, questions, answers):
#     # Create a mapping of questions to answers
#     qa_map = {q.strip(): a for q, a in zip(questions, answers)}
    
#     # Find all question paragraphs
#     question_indices, question_texts = find_question_paragraphs(doc)
    
#     modified = False
#     insertions = 0
    
#     # Process each question paragraph
#     for idx, question_text in zip(question_indices, question_texts):
#         # Find the matching answer
#         matching_answer = None
#         matching_question = None
        
#         # Try to find an exact match first
#         if question_text in qa_map:
#             matching_answer = qa_map[question_text]
#             matching_question = question_text
#         else:
#             # Try to find a fuzzy match
#             for q in qa_map:
#                 # Check if the current question is similar to any in our map
#                 if question_text in q or q in question_text:
#                     matching_answer = qa_map[q]
#                     matching_question = q
#                     break
        
#         if matching_answer:
#             # Force insert answer after the question
#             print(f"Inserting answer after: {question_text}")
            
#             # Create a new paragraph for "Answer:" heading
#             p_answer_heading = doc.add_paragraph()
#             doc._body._body.insert(idx + 1, p_answer_heading._p)
#             p_answer_heading.style = doc.styles['Normal']
#             p_answer_heading.paragraph_format.space_before = Pt(12)
#             r_answer_heading = p_answer_heading.add_run("Answer: ")
#             r_answer_heading.bold = True
#             r_answer_heading.font.color.rgb = RGBColor(0, 102, 0)
            
#             # Create a new paragraph for the answer content
#             p_answer = doc.add_paragraph()
#             doc._body._body.insert(idx + 2, p_answer._p)
#             p_answer.style = doc.styles['Normal']
#             p_answer.paragraph_format.left_indent = Pt(20)
#             r_answer = p_answer.add_run(matching_answer)
#             r_answer.italic = True
            
#             # Add a blank line after the answer
#             p_blank = doc.add_paragraph()
#             doc._body._body.insert(idx + 3, p_blank._p)
            
#             modified = True
#             insertions += 1
            
#             # Adjust subsequent indices to account for insertions
#             for j in range(len(question_indices)):
#                 if question_indices[j] > idx:
#                     question_indices[j] += 3  # We inserted 3 paragraphs
    
#     print(f"Inserted {insertions} answers into the document.")
#     return modified

# # Identify and extract all quiz questions from the document
# def extract_all_quiz_questions(doc):
#     full_text = extract_lab_manual_text(doc)
    
#     # Extract all quiz sections
#     all_questions = []
    
#     # Look for quiz sections
#     quiz_sections_pattern = r'(Quiz\(.*?\)|Quiz: \(Sufficient space to be provided for the answers\))(.*?)(References used by the students:?|Suggested Reference:?)'
#     quiz_sections = re.findall(quiz_sections_pattern, full_text, re.DOTALL)
    
#     if quiz_sections:
#         print(f"Found {len(quiz_sections)} quiz sections")
#         for i, section in enumerate(quiz_sections):
#             quiz_section = section[1].strip()  # Content between 'Quiz' and the reference section
#             questions = extract_questions_from_quiz_section(quiz_section)
#             print(f"Quiz section {i+1}: Found {len(questions)} questions")
#             all_questions.extend(questions)
#     else:
#         print("No quiz sections found. Extracting all numbered questions...")
#         # If no quiz sections, try to extract all numbered questions from the document
#         all_text = full_text
#         questions = extract_questions_from_quiz_section(all_text)
#         print(f"Found {len(questions)} numbered questions in the document")
#         all_questions.extend(questions)
    
#     return all_questions

# # Main Function
# async def main():
#     # Replace with your docx file path
#     input_file = "E:\\ClgStu\\Sem 5\\Computer Networks(3150710)\\3150710-Computer-Network Lab manual.docx"
#     output_file = "lab_manual_with_answers.docx"
    
#     print(f"Processing lab manual: {input_file}")
    
#     # Load the document
#     doc = Document(input_file)
    
#     # Extract full text for context
#     full_text = extract_lab_manual_text(doc)
    
#     # Extract all questions from the document
#     all_questions = extract_all_quiz_questions(doc)
    
#     if all_questions:
#         print(f"Found a total of {len(all_questions)} questions")
        
#         # Generate answers
#         print("Generating answers...")
#         answers = []
#         for question in all_questions:
#             answer = await generate_detailed_answer(question, full_text)
#             answers.append(answer)
        
#         # Insert answers into the document - FORCED VERSION
#         print("Inserting answers into the document (forced mode)...")
#         modified = insert_answers_in_document_forced(doc, all_questions, answers)
        
#         if modified:
#             # Save the modified document
#             doc.save(output_file)
#             print(f"Successfully created document with answers: {output_file}")
#         else:
#             print("No modifications were made to the document.")
#     else:
#         print("No questions found in the document.")
        
#     print('Process completed.')

# # Run the main function using asyncio
# if __name__ == "__main__":
#     asyncio.run(main())

import os
import re
import asyncio
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Initialize Gemini API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# Extract full text from DOCX for context
def extract_lab_manual_text(doc):
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    return "\n".join(full_text)

# Process quiz questions using Gemini API with detailed answers
async def generate_detailed_answer(question, context):
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    prompt = f"""You are an assistant that writes answers like a student writing in a college/university manual. 
    - Use clear, simple, slightly formal sentences.
    - No creative or over-explained language.
    - Keep the answer focused and to the point.
    - Format answers like textbook notes or manual entries.
    
    Format your answer for direct insertion into a Microsoft Word document:
    - Mark sections that should be bullet points with '•' at the beginning
    - Format sections that should be bold by enclosing them in <b>text to be bold</b>
    - Format sections that should be italic by enclosing them in <i>text to be italic</i>
    - Format sections that should be both bold and italic by enclosing them in <bi>text to be bold and italic</bi>
    - If you need tables, describe them in the format: <table>row1col1|row1col2;row2col1|row2col2</table>
    - For diagrams, please describe what should be in the diagram with <diagram>description</diagram>
    
    Make it look like a student textbook answer that's educational and comprehensive.
    
    Question: {question}
    
    Context: {context[:10000]}"""
    
    try:
        response = model.generate_content(prompt)
        answer = response.text.strip() if response.text else 'No answer generated.'
        print(f"\nQ: {question}\nA: {answer}\n")
        return answer
    except Exception as err:
        print(f"Error generating answer for '{question}': {str(err)}")
        return 'Error generating answer.'

# Improved Question Extraction Regex
def extract_questions_from_quiz_section(quiz_section):
    # Regex to capture numbered questions followed by any text, even if it doesn't end in a '?' or '.'.
    question_regex = r'\d+\.\s*(.*?)(?=\d+\.\s|$)'
    questions = re.findall(question_regex, quiz_section, re.DOTALL)
    return [q.strip() for q in questions]

# Function to find all paragraphs containing numbered questions
def find_question_paragraphs(doc):
    question_paragraphs = []
    question_texts = []
    question_formats = []  # Store formatting information for each question
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Match questions that start with a number followed by period, then text
        match = re.match(r'^\s*(\d+)\.\s*(.*?)$', text)
        if match:
            number = match.group(1)
            question = match.group(2).strip()
            if question:  # Only add if there's actual question text
                question_paragraphs.append(i)
                question_texts.append(question)
                
                # Store font information
                format_info = {
                    'font_name': None,
                    'font_size': None,
                    'font_color': None,
                    'alignment': None
                }
                
                # Extract format from first run that contains text
                for run in para.runs:
                    if run.text.strip():
                        if run.font.name:
                            format_info['font_name'] = run.font.name
                        if run.font.size:
                            format_info['font_size'] = run.font.size
                        if run.font.color.rgb:
                            format_info['font_color'] = run.font.color.rgb
                        break
                
                # Store paragraph alignment
                format_info['alignment'] = para.alignment
                
                question_formats.append(format_info)
                print(f"Found question {number}: {question}")
    
    return question_paragraphs, question_texts, question_formats

# Function to parse the formatted answer and apply proper Word formatting
def apply_formatting_to_paragraph(paragraph, text, format_info):
    # Process for bullet points - manually create bullet effect without using styles
    if text.startswith('•'):
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        text = text[1:].strip()
        # Add a bullet character at the beginning
        run = paragraph.add_run('• ')
        apply_base_formatting(run, format_info)
    
    # Look for formatting tags
    remaining_text = text
    position = 0
    
    # If the text starts with a bullet, we've already processed it
    if text.startswith('•'):
        remaining_text = text[1:].strip()
    
    while position < len(remaining_text):
        # Look for the next formatting tag
        bold_start = remaining_text.find('<b>', position)
        italic_start = remaining_text.find('<i>', position)
        bold_italic_start = remaining_text.find('<bi>', position)
        
        # Find the earliest tag
        next_tag_position = min(
            bold_start if bold_start != -1 else float('inf'),
            italic_start if italic_start != -1 else float('inf'),
            bold_italic_start if bold_italic_start != -1 else float('inf')
        )
        
        if next_tag_position == float('inf'):
            # No more tags, add the remaining text as regular
            if position < len(remaining_text):
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
            break
        
        # Add text before the tag
        if next_tag_position > position:
            run = paragraph.add_run(remaining_text[position:next_tag_position])
            apply_base_formatting(run, format_info)
        
        # Process the tag
        if next_tag_position == bold_start:
            end_tag = remaining_text.find('</b>', bold_start)
            if end_tag != -1:
                content = remaining_text[bold_start+3:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.bold = True
                position = end_tag + 4  # Move past </b>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
                
        elif next_tag_position == italic_start:
            end_tag = remaining_text.find('</i>', italic_start)
            if end_tag != -1:
                content = remaining_text[italic_start+3:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.italic = True
                position = end_tag + 4  # Move past </i>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
                
        elif next_tag_position == bold_italic_start:
            end_tag = remaining_text.find('</bi>', bold_italic_start)
            if end_tag != -1:
                content = remaining_text[bold_italic_start+4:end_tag]
                run = paragraph.add_run(content)
                apply_base_formatting(run, format_info)
                run.bold = True
                run.italic = True
                position = end_tag + 5  # Move past </bi>
            else:
                # Malformed tag, treat as plain text
                run = paragraph.add_run(remaining_text[position:])
                apply_base_formatting(run, format_info)
                break
    
    # Set paragraph alignment if specified
    if format_info['alignment']:
        paragraph.alignment = format_info['alignment']

# Apply base formatting (font name, size, color) to a run
def apply_base_formatting(run, format_info):
    if format_info['font_name']:
        run.font.name = format_info['font_name']
    if format_info['font_size']:
        run.font.size = format_info['font_size']
    if format_info['font_color']:
        run.font.color.rgb = format_info['font_color']

# Process table formatting
def create_table_from_description(doc, table_desc, format_info):
    # Parse table description: row1col1|row1col2;row2col1|row2col2
    rows = table_desc.split(';')
    num_rows = len(rows)
    num_cols = len(rows[0].split('|'))
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_text in enumerate(rows):
        cells = row_text.split('|')
        for j, cell_text in enumerate(cells):
            if j < num_cols:  # Ensure we don't exceed column count
                cell = table.cell(i, j)
                paragraph = cell.paragraphs[0]
                apply_formatting_to_paragraph(paragraph, cell_text.strip(), format_info)
                
    return table

# Insert answers after questions in the original document - FORCED VERSION with proper formatting
def insert_answers_in_document_forced(doc, questions, answers, question_formats):
    # Create a mapping of questions to answers
    qa_map = {q.strip(): a for q, a in zip(questions, answers)}
    
    # Find all question paragraphs with their formats
    question_indices, question_texts, question_formats = find_question_paragraphs(doc)
    
    modified = False
    insertions = 0
    
    # Process each question paragraph
    for idx, (question_text, format_info) in enumerate(zip(question_texts, question_formats)):
        question_index = question_indices[idx]
        
        # Find the matching answer
        matching_answer = None
        matching_question = None
        
        # Try to find an exact match first
        if question_text in qa_map:
            matching_answer = qa_map[question_text]
            matching_question = question_text
        else:
            # Try to find a fuzzy match
            for q in qa_map:
                # Check if the current question is similar to any in our map
                if question_text in q or q in question_text:
                    matching_answer = qa_map[q]
                    matching_question = q
                    break
        
        if matching_answer:
            # Force insert answer after the question
            print(f"Inserting answer after: {question_text}")
            
            # Create a new paragraph for "Answer:" heading
            p_answer_heading = doc.add_paragraph()
            doc._body._body.insert(question_index + 1, p_answer_heading._p)
            p_answer_heading.style = 'Normal'  # Use Normal style which should exist in every document
            p_answer_heading.paragraph_format.space_before = Pt(12)
            
            # Apply the same formatting as the question to the "Answer:" label
            r_answer_heading = p_answer_heading.add_run("Answer: ")
            apply_base_formatting(r_answer_heading, format_info)
            r_answer_heading.bold = True
            r_answer_heading.font.color.rgb = RGBColor(0, 102, 0)  # Green color for "Answer:" text
            
            # Current insertion position
            current_pos = question_index + 2
            
            # Process answer text with appropriate formatting
            answer_lines = matching_answer.split('\n')
            for line in answer_lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check for table
                table_match = re.search(r'<table>(.*?)</table>', line)
                if table_match:
                    table_desc = table_match.group(1)
                    table = create_table_from_description(doc, table_desc, format_info)
                    doc._body._body.insert(current_pos, table._tbl)
                    current_pos += 1
                    continue
                    
                # Check for diagram placeholder
                diagram_match = re.search(r'<diagram>(.*?)</diagram>', line)
                if diagram_match:
                    p_diagram = doc.add_paragraph()
                    doc._body._body.insert(current_pos, p_diagram._p)
                    r_diagram = p_diagram.add_run("[Diagram: " + diagram_match.group(1) + "]")
                    apply_base_formatting(r_diagram, format_info)
                    r_diagram.italic = True
                    current_pos += 1
                    continue
                
                # Regular paragraph with potential formatting
                p_answer = doc.add_paragraph()
                doc._body._body.insert(current_pos, p_answer._p)
                p_answer.style = 'Normal'  # Use Normal style which should exist in every document
                
                # If not a bullet point, add indentation
                if not line.startswith('•'):
                    p_answer.paragraph_format.left_indent = Pt(20)
                
                # Apply formatting to the text
                apply_formatting_to_paragraph(p_answer, line, format_info)
                current_pos += 1
            
            # Add a blank line after the answer
            p_blank = doc.add_paragraph()
            doc._body._body.insert(current_pos, p_blank._p)
            current_pos += 1
            
            modified = True
            insertions += 1
            
            # Adjust subsequent indices to account for insertions
            offset = current_pos - (question_index + 1)
            for j in range(len(question_indices)):
                if question_indices[j] > question_index:
                    question_indices[j] += offset
    
    print(f"Inserted {insertions} answers into the document.")
    return modified

# Identify and extract all quiz questions from the document
def extract_all_quiz_questions(doc):
    full_text = extract_lab_manual_text(doc)
    
    # Extract all quiz sections
    all_questions = []
    
    # Look for quiz sections
    quiz_sections_pattern = r'(Quiz\(.*?\)|Quiz: \(Sufficient space to be provided for the answers\))(.*?)(References used by the students:?|Suggested Reference:?)'
    quiz_sections = re.findall(quiz_sections_pattern, full_text, re.DOTALL)
    
    if quiz_sections:
        print(f"Found {len(quiz_sections)} quiz sections")
        for i, section in enumerate(quiz_sections):
            quiz_section = section[1].strip()  # Content between 'Quiz' and the reference section
            questions = extract_questions_from_quiz_section(quiz_section)
            print(f"Quiz section {i+1}: Found {len(questions)} questions")
            all_questions.extend(questions)
    else:
        print("No quiz sections found. Extracting all numbered questions...")
        # If no quiz sections, try to extract all numbered questions from the document
        all_text = full_text
        questions = extract_questions_from_quiz_section(all_text)
        print(f"Found {len(questions)} numbered questions in the document")
        all_questions.extend(questions)
    
    return all_questions

# Main Function
async def main():
    # Replace with your docx file path
    input_file = "E:\\ClgStu\\Sem 5\\Computer Networks(3150710)\\3150710-Computer-Network Lab manual.docx"
    output_file = "lab_manual_with_answers.docx"
    
    print(f"Processing lab manual: {input_file}")
    
    # Load the document
    doc = Document(input_file)
    
    # Extract full text for context
    full_text = extract_lab_manual_text(doc)
    
    # Extract all questions from the document
    all_questions = extract_all_quiz_questions(doc)
    
    if all_questions:
        print(f"Found a total of {len(all_questions)} questions")
        
        # Generate answers
        print("Generating answers...")
        answers = []
        for question in all_questions:
            answer = await generate_detailed_answer(question, full_text)
            answers.append(answer)
        
        # Find question paragraphs and their formats
        _, _, question_formats = find_question_paragraphs(doc)
        
        # Insert answers into the document with proper formatting
        print("Inserting answers into the document with proper formatting...")
        modified = insert_answers_in_document_forced(doc, all_questions, answers, question_formats)
        
        if modified:
            # Save the modified document
            doc.save(output_file)
            print(f"Successfully created document with answers: {output_file}")
        else:
            print("No modifications were made to the document.")
    else:
        print("No questions found in the document.")
        
    print('Process completed.')

# Run the main function using asyncio
if __name__ == "__main__":
    asyncio.run(main())